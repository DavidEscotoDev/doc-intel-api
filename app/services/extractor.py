"""Text extraction from documents."""
import pytesseract
from PIL import Image
import fitz
from docx import Document as DocxDocument
from pathlib import Path

from app.config import get_settings
from app.logging import get_logger
from app.exceptions import ProcessingError
from app.constants import MimeType, MAX_TEXT_LENGTH

logger = get_logger(__name__)


class TextExtractor:
    def __init__(self) -> None:
        self.settings = get_settings()
        self.max_text_length = MAX_TEXT_LENGTH

    def extract(self, file_path: str, mime_type: str) -> str:
        logger.info("extracting", path=file_path, mime=mime_type)

        try:
            if mime_type == MimeType.PDF:
                text = self._pdf(file_path)
            elif mime_type == MimeType.TXT:
                text = Path(file_path).read_text(encoding="utf-8", errors="replace")
            elif mime_type == MimeType.DOCX:
                text = self._docx(file_path)
            elif mime_type in (MimeType.PNG, MimeType.JPEG, MimeType.TIFF):
                text = self._image(file_path)
            else:
                raise ProcessingError(f"Unsupported MIME type: {mime_type}")

            if len(text) > self.max_text_length:
                text = text[:self.max_text_length] + "\n[TRUNCATED]"
            return text.strip()

        except ProcessingError:
            raise
        except Exception as e:
            logger.exception("extract_failed", path=file_path, error=str(e))
            raise ProcessingError(f"Extraction failed: {e}") from e

    def _pdf(self, path: str) -> str:
        parts = []
        with fitz.open(path) as doc:
            for i, page in enumerate(doc):
                t = page.get_text()
                if t.strip():
                    parts.append(f"--- Page {i+1} ---\n{t}")
        return "\n\n".join(parts)

    def _docx(self, path: str) -> str:
        doc = DocxDocument(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n\n".join(parts)

    def _image(self, path: str) -> str:
        try:
            with Image.open(path) as img:
                if img.mode != "RGB":
                    img = img.convert("RGB")
                return pytesseract.image_to_string(img, lang=self.settings.processing.ocr_language)
        except Exception as e:
            raise ProcessingError(f"OCR failed: {e}") from e


_extractor: TextExtractor | None = None


def get_text_extractor() -> TextExtractor:
    global _extractor
    if _extractor is None:
        _extractor = TextExtractor()
    return _extractor