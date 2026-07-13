"""Text extractor tests."""
import pytest
from unittest.mock import patch, MagicMock
import tempfile
import os

from app.services.extractor import TextExtractor, get_text_extractor
from app.exceptions import ProcessingError
from app.constants import MimeType


class TestTextExtractor:
    @pytest.fixture
    def extractor(self):
        return TextExtractor()

    def test_extract_txt(self, extractor):
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
            f.write("Hello, World!\nTest line 2.")
            path = f.name
        try:
            text = extractor.extract(path, MimeType.TXT)
            assert "Hello, World!" in text
            assert "Test line 2" in text
        finally:
            os.unlink(path)

    def test_extract_pdf(self, extractor):
        import fitz
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            path = f.name
        try:
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((50, 50), "PDF Content\nLine 2")
            doc.save(path)
            doc.close()
            text = extractor.extract(path, MimeType.PDF)
            assert "PDF Content" in text
        finally:
            os.unlink(path)

    def test_extract_docx(self, extractor):
        from docx import Document as DocxDocument
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            doc = DocxDocument()
            doc.add_paragraph("DOCX Content")
            doc.add_paragraph("Second para")
            doc.save(path)
            text = extractor.extract(path, MimeType.DOCX)
            assert "DOCX Content" in text
            assert "Second para" in text
        finally:
            os.unlink(path)

    @patch("app.services.extractor.pytesseract.image_to_string")
    def test_extract_image(self, mock_ocr, extractor):
        mock_ocr.return_value = "OCR text"
        from PIL import Image
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
            path = f.name
        try:
            img = Image.new("RGB", (100, 100), "white")
            img.save(path)
            text = extractor.extract(path, MimeType.PNG)
            assert text == "OCR text"
        finally:
            os.unlink(path)

    def test_unsupported_type(self, extractor):
        with tempfile.NamedTemporaryFile(suffix=".xyz", delete=False) as f:
            f.write(b"test")
            path = f.name
        try:
            with pytest.raises(ProcessingError) as exc:
                extractor.extract(path, "application/unknown")
            assert "Unsupported MIME type" in str(exc.value)
        finally:
            os.unlink(path)

    def test_truncates_long_text(self, extractor):
        long_text = "x" * 150000
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
            f.write(long_text)
            path = f.name
        try:
            text = extractor.extract(path, MimeType.TXT)
            assert len(text) <= extractor.max_text_length + 20
            assert "[TRUNCATED]" in text
        finally:
            os.unlink(path)


class TestExtractorFactory:
    def test_singleton(self):
        e1 = get_text_extractor()
        e2 = get_text_extractor()
        assert e1 is e2